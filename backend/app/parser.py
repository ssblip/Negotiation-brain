"""
Quote document parser.
Accepts PDF, Excel (.xlsx/.xls), Word (.docx), or plain text.
Uses Claude to extract structured vendor data from unstructured content.
"""
from __future__ import annotations

import io
import json

import anthropic
import pdfplumber
import openpyxl
import docx as python_docx

from app.config import settings

_client = anthropic.Anthropic(api_key=settings.anthropic_api_key)


def extract_text_from_file(filename: str, content: bytes) -> str:
    """Extract raw text from uploaded file bytes."""
    fn = filename.lower()

    if fn.endswith(".pdf"):
        text_parts: list[str] = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        return "\n".join(text_parts)

    if fn.endswith((".xlsx", ".xls")):
        wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
        rows: list[str] = []
        for sheet in wb.worksheets:
            rows.append(f"=== Sheet: {sheet.title} ===")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(c.strip() for c in cells):
                    rows.append("\t".join(cells))
        return "\n".join(rows)

    if fn.endswith(".docx"):
        doc = python_docx.Document(io.BytesIO(content))
        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text.strip() for cell in row.cells))
        return "\n".join(parts)

    # Plain text fallback
    return content.decode("utf-8", errors="replace")


_EXTRACT_PROMPT = """\
You are a procurement data extraction assistant. Extract structured vendor data from the quote \
document below and return a JSON array — one object per vendor.

Each vendor object must have these top-level fields (null if not found):
  vendor_company, vendor_name, vendor_email,
  quoted_price (number), quoted_currency (default "USD"),
  quoted_delivery_days (integer), quoted_payment_days (integer, e.g. Net-30 → 30),
  quoted_warranty_months (integer),
  raw_quote_text (the vendor's full quote section),
  custom_spec_values (object — extract ALL fields listed below)

━━━ EXTRACTION RULES FOR custom_spec_values ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
• Do NOT require exact label matches. Semantically search the vendor's text — a value
  may appear in a sentence, bullet list, table, or parenthetical.
• Use null only when the information is genuinely absent.
• Use the EXACT key names shown (quoted, case-sensitive).
• For specs that appear inside a comma-separated list (e.g. certifications), extract
  the specific relevant value rather than copying the whole list unless told otherwise.

{spec_extraction_block}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Return ONLY the JSON array. No markdown fences, no explanation.

DOCUMENT:
{text}
"""

# Standard fields always extracted into custom_spec_values
_STANDARD_SPECS: list[tuple[str, str]] = [
    ("certifications",   "string — ALL certifications, compliance codes and ratings found anywhere in the text (e.g. \"ISO 9001, MIL-STD-810H, IP65, ATEX Zone 2\")"),
    ("defect_rate",      "string — defect rate, AQL level or rejection rate if mentioned (e.g. \"0.5%\")"),
    ("quality_standard", "string — quality management framework (e.g. \"Six Sigma\", \"AQL 1.0\")"),
    ("inspection_terms", "string — inspection, FAT or acceptance test terms (e.g. \"third-party QC accepted\")"),
    ("quality_notes",    "string — any other quality claims, SLA statements or deflection language"),
    ("lead_time_risk",   "string — any delivery caveats or risk statements (e.g. \"subject to freight delays\")"),
    ("advance_payment",  "string — upfront or advance payment requirement (e.g. \"30% on PO\")"),
    ("volume_discount",  "string — volume discount terms (e.g. \"3% discount above 250 units\")"),
    ("origin_country",   "string — country of manufacture or origin"),
    ("notes",            "string — any other notable terms not captured above"),
]

_FIELD_TYPE_HINTS: dict[str, str] = {
    "NUM":    "extract as a number",
    "PCTNUM": "extract as a percentage number",
    "BOOL":   "extract as true or false (yes/no/present/absent all count)",
    "CAT":    "extract the exact value/code string — search semantically, not by label",
    "MULTI":  "extract as a comma-separated list of all matching values found",
    "TIER":   "extract the tier or level string",
    "TEXT":   "extract the relevant text snippet",
}


def _build_spec_extraction_block(custom_specs: list[dict]) -> str:
    lines: list[str] = []
    n = 1

    # Standard fields — always present
    lines.append("STANDARD FIELDS (always extract these):")
    for key, desc in _STANDARD_SPECS:
        lines.append(f'  {n}. "{key}": {desc}')
        n += 1

    # Buyer-defined specs — dynamic, from Step 1
    if custom_specs:
        lines.append("")
        lines.append("BUYER-DEFINED SPECS (defined by the buyer for this negotiation):")
        for s in custom_specs:
            name = s["name"]
            ft   = s.get("field_type", "TEXT")
            req  = s.get("required_value")
            unit = f" [{s['unit']}]" if s.get("unit") else ""
            hint = _FIELD_TYPE_HINTS.get(ft, "extract the value")
            example = f' — target value is "{req}"' if req else ""
            lines.append(f'  {n}. "{name}"{unit}: {hint}{example}')
            n += 1
        lines.append("")
        lines.append('  Example: "IP Rating" type CAT, vendor text "Certifications: MIL-STD-810H, IP65"')
        lines.append('  → custom_spec_values["IP Rating"] = "IP65"  (extracted from the list semantically)')

    return "\n".join(lines)


_CONDENSE_PROMPT = """\
You are extracting chatbot operating instructions from a procurement negotiation strategy document.
Output ONLY the content the chatbot needs at runtime. Compact bullet points. No preamble, no commentary.
Target: under 1000 words total.

Extract the following sections (use these exact headings). Omit any section absent from the document.

## CORE PRINCIPLES
Up to 6 bullets. The non-negotiable rules the bot must always follow regardless of strategy.

## STRATEGIES
For EACH strategy variant found (e.g. S1–S6): one entry with:
  - Name & trigger condition (when to use this strategy)
  - Primary tactic / opening stance
  - Key moves and responses
  - What to concede first vs last
  - Exit / escalation condition for this strategy

## PRICE NEGOTIATION RULES
All rules about price moves, including:
  - Maximum single-round concession % allowed before escalating (CRITICAL — preserve exact threshold and action)
  - How to respond to price anchors, large jumps, and aggressive cuts
  - When to invoke "best and final offer" (BAFO)
  - When price is off-limits vs negotiable

## SPEC / QUALITY / DELIVERY / PAYMENT / WARRANTY RULES
Key rules per dimension:
  - How to handle quality deflection (vendor claims spec changes solve price gap)
  - Delivery negotiation approach and limits
  - Payment terms approach and limits
  - Warranty approach
  - Logrolling: how to trade one dimension against another

## CONCESSION STRATEGY
  - Concession pacing rules (how fast / how much per round)
  - Diminishing concession pattern if specified
  - What to give up first vs protect longest
  - How to signal "near limit" without revealing the actual limit

## ESCALATION TRIGGERS
Exact conditions that require escalating to a human buyer, including:
  - Large single-round price move threshold (preserve exact % and wording)
  - Red-flag vendor tactics that trigger escalation
  - Any other escalation conditions listed

## BATNA & BAFO PROTOCOLS
  - When to invoke BATNA (best alternative)
  - How to present BAFO
  - How to close or walk away

## BEHAVIORAL SCENARIOS
  - How to respond to each vendor tactic listed (e.g. anchoring, urgency pressure, quality deflection, lowball, bundling tricks)
  - Red flags to watch for

## FORBIDDEN & PERMITTED LANGUAGE
  - Exact phrases or types of statements the bot MUST NEVER say (forbidden)
  - Exact phrases or approaches the bot SHOULD use (permitted/preferred)
  - Secrecy rules: what internal numbers / targets must never be revealed

## AGREEMENT & HANDOFF
  - What constitutes a valid agreement
  - What the bot says / does not say after agreement is reached
  - No award language rules

DOCUMENT:
{text}
"""


def condense_strategy_doc(raw_text: str) -> str:
    """Run strategy doc through Claude to extract only chatbot-relevant logic (~500-800 tokens)."""
    prompt = _CONDENSE_PROMPT.format(text=raw_text[:40000])
    msg = _client.messages.create(
        model=settings.claude_model,
        max_tokens=1500,
        messages=[{"role": "user", "content": prompt}],
    )
    return msg.content[0].text.strip()


def parse_vendors_from_text(raw_text: str, custom_specs: list[dict] | None = None) -> list[dict]:
    """Use Claude to extract vendor quote data from raw document text."""
    spec_block = _build_spec_extraction_block(custom_specs or [])
    prompt = _EXTRACT_PROMPT.format(
        text=raw_text[:30000],
        spec_extraction_block=spec_block,
    )

    msg = _client.messages.create(
        model=settings.claude_model,
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}],
    )
    response_text = msg.content[0].text.strip()

    # Strip accidental markdown fences
    if response_text.startswith("```"):
        lines = response_text.split("\n")
        response_text = "\n".join(lines[1:-1])

    try:
        vendors = json.loads(response_text)
        if isinstance(vendors, dict):
            vendors = [vendors]
        return vendors
    except json.JSONDecodeError:
        return []
